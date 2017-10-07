import csv
import docx2txt
import os
import re
import subprocess


from docx import Document
from pathlib import Path
from shutil import copyfile

from . helpers import DateRange

# For personal use in linux/ubuntu
# LibreOffice is required.
try:
  subprocess.call(['soffice', '--help'], stdout=subprocess.PIPE)
except FileNotFoundError:
  print("LibreOffice is required, but not installed.")
  exit()

class Form:
  def __init__(self, path):
    """
    Arguments:
      path (string) : Absolute, or relative to cwd, path to file.
    """

    # LibreOffice is needed for conversion but can be run only once.
    # Make sure it is closed.
    result = str(subprocess.check_output(('ps', '-A')))
    while 'soffice' in result:
      print('\n' * 80)
      print("Please close LibreOffice first to ensure this program works correctly.")
      input('Press enter to continue.')
      result = str(subprocess.check_output(('ps', '-A')))

    # Attributes
    self.path = path
    self.directory, self.full_name = os.path.split(self.path)
    self.parentdir = os.path.split(str(Path(path).parents[0]))[1]
    self.short_name, self.extension = os.path.splitext(self.full_name)
    self.plant_name = ''
    self.plant_part = ''
    self.date = ''

    # Allowed file types
    if self.extension.lower() not in ('.doc', '.docx'):
      raise TypeError("File format {} not accepted.".format(self.extension))

    # Convert doc to docx, if it has not already been done.
    if self.extension.lower() == '.doc':
      new_path = os.path.join('tmp', self.parentdir, self.short_name + '.docx')
      # Does converted file already exist?
      try:
        open(new_path)
      # Convert file if .docx version does not yet exist.
      except FileNotFoundError:
        self.convert_to_docx()
      finally:
        self.path = new_path
      self.doc = Document(self.path)

  def convert_to_txt(self):
    """Save a copy of form as txt file"""
    try:
      open(os.path.join('tmp', self.directory, self.short_name + '.txt'))
    except FileNotFoundError:
      args = [
        'soffice',
        '--headless',
        '--convert-to',
        'txt:Text',
        '--outdir',
        'tmp/' + self.directory,
        self.path,
      ]
      subprocess.call(args)

  def get_plant_name(self):
    """Get plant name from form."""
    if not self.plant_name:
      self.plant_name = self._search('botanische naam')
    return self.plant_name

  def get_plant_part(self):
    """Get harvested plant part from form."""
    if not self.plant_part:
      self.plant_part = self._search('Geleverde plantendeel')
    return self.plant_part

  def get_date(self):
    """ Get date from form.
    Accepted date format:
      (d)d-(m)m-(yy)yy
      (d) month (yy)hh
    Dates may be separated by comma's if multiple dates are documented.
    """
    # Collect rough date from form
    if not self.date:
      rough_date = self._search('Leverdatum')
      if rough_date:
        # Are there multiple dates separated by a comma?
        if ',' in rough_date:
          dates = rough_date.strip(', ').split(',')
          dates = [self._read_date(d) for d in dates]
          # Get median
          self.date = dates[len(dates)//2]
        # Only one date
        else:
          self.date = self._read_date(rough_date)
    return self.date

  def _read_date(self, date):
    """ Format string as dd-mm-yyyy string. """
    # Is month notated as letters? eg. dd-month-yyyy
    pattern = re.compile(r'(\d+).*?([a-zA-Z]+).*?(\d+)')
    month_as_string = True
    # Is month notated as numbers? e.g. dd-mm-yyyyy
    if not pattern.search(date):
      pattern = re.compile(r'(\d+).*-.*?(\d+).*-.*?(\d+)')
      month_as_string = False

    match = pattern.search(date)
    if match:
      day, month, year = [str(g) for g in match.groups()]
      # Pad 2 digit years with 20
      year = year if len(year) > 2 else '20' + year
      # Convert month to a number if necessary
      if month_as_string:
        month = str(DateRange.months.index(month)).rjust(2, '0')
    # If no match is found
    else:
      day, month, year = ('00', '00', '0000')
    # Return the date, add 0's where needed, to ensure consistent formatting.
    return day.rjust(2, '0') + '-' + month.rjust(2, '0') + '-' + year

  def _search(self, query):
    """ Search document for cell.
    Return the value from the cell next to it. """
    for table in self.doc.tables:
      for row in table.rows:
        i = 0
        for cell in row.cells:
          if query.strip().lower() in cell.text.lower():
            c = row.cells[i+1]
            return c.text
          i += 1

  def _index_of_cell(self, query):
    """ Search document for cell.
    Return the index of the cell (table, row, cell)
    """
    x = 0
    for table in self.doc.tables:
      y = 0
      for row in table.rows:
        z = 0
        for cell in row.cells:
          if query.strip().lower() in cell.text.lower():
            return (x, y, z)
          z += 1
        y += 1
      x += 1

  def convert_to_docx(self):
    """ Save a copy of form as docx. Return new path. """
    args = ['soffice', '--headless', '--convert-to', 'docx',
            '--outdir', 'tmp/' + self.parentdir, self.path]
    subprocess.call(args, stdout=subprocess.PIPE)
    return 'tmp/' + self.path + 'x'


class Spec(Form):
  """ SPC form. """
  not_found = []
  not_enough_data = []

  def __init__(self, path):
    # Call parent init
    super().__init__(path)

    # Changed SPC files will be saved to output/SPC
    if not os.path.exists('output/SPC'):
      os.mkdir('output/SPC')
    # Copy file to output/SPC if it has not already been done.
    new_path = os.path.join('output', 'SPC', self.short_name + '.docx')
    if not os.path.isfile(new_path):
      copyfile(self.path, new_path)
    # Change working file to the file in the output/SPC directory.
    self.path = new_path

  def write_name_and_ref_to_csv(self, filename="names_and_refs", delim=","):
    """Save the plant name and reference number to file """
    extension = '.csv'
    filename = filename + extension
    with open(filename, 'a') as f:
      f.write('{},{}\n'.format(self.plant_name, self.refname))

  def update(self, data=None, min_n=None):
    """Update spec based on data.
    Arguments:
      min_n (int): Minimum number of data entries required to update date/range. TODO
    """
    # Collect data
    if data is not None:
      self.data = []
      with open(data) as f:
        for row in csv.reader(f):
          self.data.append(row)
    else:
      self.data = Data().get_data()

    # Open file
    self.doc = Document(self.path)
    return self._edit_cells(min_n)

  def _edit_cells(self, min_n):
    """ Edit the cells in the SPC form file. """
    # Get scientific name
    self.plant_name = self.get_plant_name()
    if not self.plant_name:
      self.not_found.append(self.full_name)
      return

    # Get plant part
    # The plant part appears in the header which cannot as of yet be read by docx
    # docx2txt is used instead.
    text = docx2txt.process(self.path)
    self.plant_part = text.split('\n', 6)[6].split(',')[1] # the location of the plant part
    if not self.plant_part:
      self.not_found.append(self.full_name)
      return

    # Get coordinates of cell containing period
    x, y, z = self._index_of_cell('oogstperiode')
    c = self.doc.tables[x].rows[y].cells[z+1]
    old_period = c.text

    # Search for harvest period in data
    new_period = ''
    for line in self.data:
      # Get first two words of full name 
      # which appears in the second column in the list.
      name = ' '.join(line[0].strip('"').split()[:2])
      part = line[1]
      n_data = int(line[4]) # n harvests data is based on

      # Check if there is sufficient data for plant/part
      if min_n:
        if n_data < min_n:
          self.not_enough_data.append(self.full_name)
          return

      # Compare name in data to name found in document
      if name.lower() in self.plant_name.lower():
        # Compare part in data to part in spec
        if part.lower()[:4] in self.plant_part.lower():
          # Harvest period is in the 3rd column
          new_period= line[2]
          break

    if not new_period:
      self.not_found.append(self.full_name + ' ' + self.plant_name)
      return

    # Edit file
    self.doc.styles['Normal'].font.name = 'Arial'
    c.text = new_period
    self.doc.save(self.path)

    # Success
    return (self.full_name, self.plant_name, self.plant_part.lower().strip(), old_period, new_period, n_data)

